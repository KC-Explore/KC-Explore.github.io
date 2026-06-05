#!/usr/bin/env python3
"""
build_docx.py — render the Jekyll Markdown posts into Word .docx review artifacts.

The Markdown in _posts/ is the single source of truth. This script never edits it;
it only produces ./docx/*.docx (one per post) plus a combined master document so
Keval can review/comment in Word before anything is published.

Preferred renderer is pandoc (best fidelity) if it's on PATH; otherwise this falls
back to a self-contained python-docx renderer that handles the subset of Markdown
used in these posts: front matter, headings, paragraphs, bold/italic/inline-code,
links, bullet & numbered lists, fenced code blocks, blockquotes, tables, hr, and
the <div class="sidebar-note"|"callout"> callout blocks.

Usage:
    python3 build_docx.py
"""

import os
import re
import glob
import shutil
import subprocess

ROOT = os.path.dirname(os.path.abspath(__file__))
POSTS_DIR = os.path.join(ROOT, "_posts")
OUT_DIR = os.path.join(ROOT, "docx")


# ----------------------------------------------------------------------------- #
# Front matter
# ----------------------------------------------------------------------------- #
def split_front_matter(text):
    """Return (meta_dict, body) — strips a leading --- ... --- YAML-ish block."""
    meta = {}
    if text.startswith("---"):
        end = text.find("\n---", 3)
        if end != -1:
            fm = text[3:end].strip()
            body = text[end + 4:].lstrip("\n")
            for line in fm.splitlines():
                if ":" in line:
                    k, _, v = line.partition(":")
                    meta[k.strip()] = v.strip().strip('"').strip("'")
            return meta, body
    return meta, text


def strip_callout_divs(body):
    """Turn <div class="sidebar-note">...</div> wrappers into a marker line so the
    fallback renderer can style them, and drop the raw HTML tags."""
    body = re.sub(
        r'<div class="sidebar-note"[^>]*>',
        "\n[[SIDEBAR]]\n",
        body,
    )
    body = re.sub(r'<div class="callout"[^>]*>', "\n[[CALLOUT]]\n", body)
    body = body.replace("</div>", "\n[[ENDBOX]]\n")
    return body


# ----------------------------------------------------------------------------- #
# pandoc path
# ----------------------------------------------------------------------------- #
def build_with_pandoc(md_path, docx_path):
    # Strip front matter to a temp file so pandoc doesn't render YAML as a table.
    with open(md_path, encoding="utf-8") as f:
        _, body = split_front_matter(f.read())
    tmp = docx_path + ".tmp.md"
    with open(tmp, "w", encoding="utf-8") as f:
        f.write(body)
    subprocess.run(
        ["pandoc", tmp, "-f", "gfm", "-o", docx_path], check=True
    )
    os.remove(tmp)


# ----------------------------------------------------------------------------- #
# python-docx fallback renderer
# ----------------------------------------------------------------------------- #
INLINE_RE = re.compile(
    r"(\*\*.+?\*\*|\*.+?\*|`.+?`|\[.+?\]\(.+?\))"
)


def add_inline(paragraph, text):
    """Render **bold**, *italic*, `code`, and [links] into runs."""
    for part in INLINE_RE.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            paragraph.add_run(part[2:-2]).bold = True
        elif part.startswith("*") and part.endswith("*"):
            paragraph.add_run(part[1:-1]).italic = True
        elif part.startswith("`") and part.endswith("`"):
            r = paragraph.add_run(part[1:-1])
            r.font.name = "Consolas"
        elif part.startswith("[") and "](" in part:
            label = part[1:part.index("]")]
            url = part[part.index("](") + 2:-1]
            r = paragraph.add_run(f"{label} ({url})")
            r.italic = True
        else:
            paragraph.add_run(part)


def render_fallback(doc, body):
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    lines = body.splitlines()
    i = 0
    n = len(lines)
    while i < n:
        line = lines[i]
        stripped = line.strip()

        # Fenced code block
        if stripped.startswith("```"):
            i += 1
            code = []
            while i < n and not lines[i].strip().startswith("```"):
                code.append(lines[i])
                i += 1
            i += 1  # skip closing fence
            p = doc.add_paragraph()
            r = p.add_run("\n".join(code))
            r.font.name = "Consolas"
            r.font.size = Pt(9)
            return_shade(p)
            continue

        # Callout markers
        if stripped in ("[[SIDEBAR]]", "[[CALLOUT]]"):
            label = "How I run mine" if stripped == "[[SIDEBAR]]" else "Heads up"
            p = doc.add_paragraph()
            r = p.add_run(f"  {label}  ")
            r.bold = True
            r.font.color.rgb = RGBColor(0x2F, 0x6B, 0xD8)
            i += 1
            continue
        if stripped == "[[ENDBOX]]":
            i += 1
            continue

        # Horizontal rule
        if re.match(r"^(-{3,}|\*{3,}|_{3,})$", stripped):
            doc.add_paragraph().add_run("—" * 20)
            i += 1
            continue

        # Headings
        m = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if m:
            level = len(m.group(1))
            doc.add_heading(m.group(2).strip(), level=min(level, 4))
            i += 1
            continue

        # Tables
        if "|" in line and i + 1 < n and re.match(r"^\s*\|?[\s:|-]+\|?\s*$", lines[i + 1]):
            header = [c.strip() for c in stripped.strip("|").split("|")]
            i += 2  # header + separator
            rows = []
            while i < n and "|" in lines[i] and lines[i].strip():
                rows.append([c.strip() for c in lines[i].strip().strip("|").split("|")])
                i += 1
            table = doc.add_table(rows=1, cols=len(header))
            table.style = "Light Grid Accent 1"
            for j, h in enumerate(header):
                add_inline(table.rows[0].cells[j].paragraphs[0], h)
            for row in rows:
                cells = table.add_row().cells
                for j, c in enumerate(row[: len(header)]):
                    add_inline(cells[j].paragraphs[0], c)
            continue

        # Blockquote
        if stripped.startswith(">"):
            p = doc.add_paragraph(style="Intense Quote")
            add_inline(p, stripped.lstrip("> ").strip())
            i += 1
            continue

        # Bullet list
        if re.match(r"^[-*]\s+", stripped):
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, re.sub(r"^[-*]\s+", "", stripped))
            i += 1
            continue

        # Numbered list
        if re.match(r"^\d+\.\s+", stripped):
            p = doc.add_paragraph(style="List Number")
            add_inline(p, re.sub(r"^\d+\.\s+", "", stripped))
            i += 1
            continue

        # Blank line
        if not stripped:
            i += 1
            continue

        # Paragraph
        p = doc.add_paragraph()
        add_inline(p, stripped)
        i += 1


def return_shade(paragraph):
    """Light grey shading behind a code paragraph."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), "F0F1F3")
    pPr.append(shd)


def build_with_python_docx(md_path, docx_path):
    from docx import Document

    with open(md_path, encoding="utf-8") as f:
        meta, body = split_front_matter(f.read())
    body = strip_callout_divs(body)

    doc = Document()
    title = meta.get("title", os.path.basename(md_path))
    doc.add_heading(title, level=0)
    render_fallback(doc, body)
    doc.save(docx_path)


# ----------------------------------------------------------------------------- #
# main
# ----------------------------------------------------------------------------- #
def main():
    os.makedirs(OUT_DIR, exist_ok=True)
    use_pandoc = shutil.which("pandoc") is not None
    print(f"Renderer: {'pandoc' if use_pandoc else 'python-docx fallback'}")

    posts = sorted(glob.glob(os.path.join(POSTS_DIR, "*.md")))
    if not posts:
        print("No posts found in _posts/ — nothing to build.")
        return

    for md in posts:
        base = os.path.splitext(os.path.basename(md))[0]
        out = os.path.join(OUT_DIR, base + ".docx")
        if use_pandoc:
            build_with_pandoc(md, out)
        else:
            build_with_python_docx(md, out)
        print(f"  -> {os.path.relpath(out, ROOT)}")

    # Combined master document (fallback renderer only; pandoc can take all files).
    master = os.path.join(OUT_DIR, "00-MASTER-all-posts.docx")
    if use_pandoc:
        bodies = []
        for md in posts:
            with open(md, encoding="utf-8") as f:
                _, b = split_front_matter(f.read())
            bodies.append(b)
        tmp = master + ".tmp.md"
        with open(tmp, "w", encoding="utf-8") as f:
            f.write("\n\n\\newpage\n\n".join(bodies))
        subprocess.run(["pandoc", tmp, "-f", "gfm", "-o", master], check=True)
        os.remove(tmp)
    else:
        from docx import Document

        doc = Document()
        doc.add_heading("Start Your AI Homelab — Full Series", level=0)
        for md in posts:
            with open(md, encoding="utf-8") as f:
                meta, body = split_front_matter(f.read())
            body = strip_callout_divs(body)
            doc.add_heading(meta.get("title", os.path.basename(md)), level=1)
            render_fallback(doc, body)
            doc.add_page_break()
        doc.save(master)
    print(f"  -> {os.path.relpath(master, ROOT)}")
    print("Done.")


if __name__ == "__main__":
    main()
